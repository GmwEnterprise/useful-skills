#!/usr/bin/env python3
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}

OLE_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
)
IMAGE_REL_TYPES = {
    "http://schemas.openxmlformats.org/drawingml/2006/relationship/image",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
}


def format_paragraph_text(paragraph) -> str:
    parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        bold = run.bold
        italic = run.italic
        strike = run.font.strike if run.font else None
        underline = run.underline

        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"

        if strike:
            text = f"~~{text}~~"
        if underline:
            text = f"<u>{text}</u>"

        parts.append(text)

    return "".join(parts)


def heading_level(paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.replace("Heading", "").strip())
        except ValueError:
            return None
    return None


def table_to_markdown(table) -> str:
    if not table.rows:
        return ""

    lines = []
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
            cells.append(cell_text if cell_text else " ")
        lines.append("| " + " | ".join(cells) + " |")

        if row_idx == 0:
            lines.append("| " + " | ".join("---" for _ in cells) + " |")

    return "\n".join(lines)


def extract_images_and_attachments(docx_path: str, assets_dir: Path) -> list[str]:
    extracted = []
    docx_file = Path(docx_path)
    assets_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(docx_file, "r") as zf:
        attachment_extensions = {
            ".xlsx",
            ".xls",
            ".pptx",
            ".ppt",
            ".pdf",
            ".doc",
            ".docx",
            ".zip",
            ".rar",
            ".csv",
            ".txt",
        }

        for entry in zf.namelist():
            name = Path(entry)
            ext = name.suffix.lower()

            if name.parent.name == "media" and "word" in entry.split("/"):
                out_path = assets_dir / name.name
            elif ext in attachment_extensions and "embeddings" in entry.split("/"):
                out_path = assets_dir / name.name
            else:
                continue

            if out_path.exists():
                stem = out_path.stem
                counter = 1
                while out_path.exists():
                    out_path = assets_dir / f"{stem}_{counter}{ext}"
                    counter += 1

            data = zf.read(entry)
            out_path.write_bytes(data)
            extracted.append(str(out_path))

    return extracted


def build_image_map(docx_path: str) -> dict[str, str]:
    image_map: dict[str, str] = {}
    docx_file = Path(docx_path)

    with zipfile.ZipFile(docx_file, "r") as zf:
        if "word/_rels/document.xml.rels" not in zf.namelist():
            return image_map

        rels_xml = zf.read("word/_rels/document.xml.rels")
        root = ET.fromstring(rels_xml)

        for rel in root:
            rid = rel.get("Id", "")
            target = rel.get("Target", "")
            rel_type = rel.get("Type", "")

            if any(t in rel_type for t in ("image", "oleObject")):
                if target.startswith("/"):
                    target = target.lstrip("/")
                elif not target.startswith("word/"):
                    target = "word/" + target

                target_name = Path(target).name
                image_map[rid] = target_name

    return image_map


def get_inline_images(
    paragraph, image_map: dict[str, str], assets_name: str
) -> list[str]:
    refs = []
    p_elem = paragraph._element

    for drawing in p_elem.findall(".//w:drawing", NS):
        blip = drawing.find(".//a:blip", NS)
        if blip is None:
            for imagedata in drawing.findall(".//", NS):
                src = imagedata.get("src", "")
                if src and src in image_map:
                    filename = image_map[src]
                    refs.append(f"![{filename}]({assets_name}/{filename})")
            continue

        rid = blip.get(f"{{{NS['r']}}}embed", "")
        if rid and rid in image_map:
            filename = image_map[rid]
            refs.append(f"![{filename}]({assets_name}/{filename})")

    return refs


def extract_metadata(doc: Document, docx_path: str) -> str:
    props = doc.core_properties
    lines = ["---"]

    if props.title:
        lines.append(f'title: "{props.title}"')
    if props.author:
        lines.append(f'author: "{props.author}"')
    if props.created:
        lines.append(f'created: "{props.created.isoformat()}"')
    if props.modified:
        lines.append(f'modified: "{props.modified.isoformat()}"')
    if props.subject:
        lines.append(f'subject: "{props.subject}"')
    if props.keywords:
        lines.append(f'keywords: "{props.keywords}"')
    if props.category:
        lines.append(f'category: "{props.category}"')

    lines.append(f'source: "{Path(docx_path).name}"')
    lines.append("---")

    return "\n".join(lines)


def docx_to_markdown(docx_path: str, output_dir: Path) -> str:
    doc = Document(docx_path)
    docx_name = Path(docx_path).stem
    assets_name = docx_name
    assets_dir = output_dir / assets_name

    extract_images_and_attachments(docx_path, assets_dir)
    image_map = build_image_map(docx_path)

    md_lines = []

    metadata = extract_metadata(doc, docx_path)
    md_lines.append(metadata)
    md_lines.append("")

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    md_lines.append(table_to_markdown(table))
                    md_lines.append("")
                    break
            continue

        if tag != "p":
            continue

        for paragraph in doc.paragraphs:
            if paragraph._element is not element:
                continue

            text = paragraph.text.strip()
            fmt_text = format_paragraph_text(paragraph)

            inline_images = get_inline_images(paragraph, image_map, assets_name)

            level = heading_level(paragraph)
            if level is not None:
                md_lines.append(f"{'#' * level} {fmt_text}")
                md_lines.append("")
            elif paragraph.style and "List" in (paragraph.style.name or ""):
                prefix = "- "
                if fmt_text.startswith(prefix):
                    md_lines.append(fmt_text)
                else:
                    md_lines.append(f"- {fmt_text}")
                md_lines.append("")
            elif not text and inline_images:
                for img in inline_images:
                    md_lines.append(img)
                    md_lines.append("")
            elif not text:
                md_lines.append("")
            else:
                alignment = paragraph.alignment
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    md_lines.append(f'<div align="center">{fmt_text}</div>')
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    md_lines.append(f'<div align="right">{fmt_text}</div>')
                else:
                    md_lines.append(fmt_text)
                md_lines.append("")

            if text and inline_images:
                for img in inline_images:
                    md_lines.append(img)
                    md_lines.append("")

            break

    result = "\n".join(md_lines)
    while "\n\n\n" in result:
        result = result.replace("\n\n\n", "\n\n")

    return result.strip() + "\n"


def main():
    if len(sys.argv) < 2:
        print("Usage: python main.py <docx_file> [output_dir]", file=sys.stderr)
        print("\nArguments:", file=sys.stderr)
        print("  docx_file   - Path to the .docx file", file=sys.stderr)
        print(
            "  output_dir  - Optional: Output directory (default: ./docs/extracted)",
            file=sys.stderr,
        )
        sys.exit(1)

    file_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        path = Path(file_path)
        if not path.exists():
            print(f"Error: File not found: {file_path}", file=sys.stderr)
            sys.exit(1)

        if path.suffix.lower() != ".docx":
            print(
                f"Error: Unsupported format: {path.suffix}. Only .docx is supported.",
                file=sys.stderr,
            )
            sys.exit(1)

        out = Path(output_dir) if output_dir else Path.cwd() / "docs" / "extracted"

        out.mkdir(parents=True, exist_ok=True)

        md_content = docx_to_markdown(str(path), out)

        docx_name = path.stem
        md_file = out / f"{docx_name}.md"

        md_file.write_text(md_content, encoding="utf-8")

        assets_dir = out / docx_name
        asset_count = 0
        if assets_dir.exists():
            asset_count = len([f for f in assets_dir.iterdir() if f.is_file()])

        print(f"Success: {path.name}")
        print(f"Markdown: {md_file}")
        if asset_count > 0:
            print(f"Assets: {assets_dir} ({asset_count} files)")

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
